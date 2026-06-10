from __future__ import annotations

from app.core.logging import get_logger
from app.domain.models import FileDescriptor, ParsedDocument, ParsedSection
from app.services.parsers.assets import write_image_asset
from app.services.parsers.base import BaseParser

logger = get_logger(component="docx_parser")


class DOCXParser(BaseParser):
    def parse(self, file_descriptor: FileDescriptor) -> ParsedDocument:
        from docx import Document

        doc = Document(str(file_descriptor.file_path))

        sections: list[ParsedSection] = []
        current_heading = ""
        buffer: list[str] = []

        def flush_buffer() -> None:
            if not buffer:
                return
            text = "\n".join(buffer).strip()
            if text:
                sections.append(ParsedSection(text=text, section_title=current_heading))
            buffer.clear()

        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text.strip()
            if not paragraph_text:
                continue

            style_name = (paragraph.style.name or "").lower()
            if style_name.startswith("heading"):
                flush_buffer()
                current_heading = paragraph_text
                continue

            buffer.append(paragraph_text)

        flush_buffer()

        table_sections = self._extract_tables(doc)
        sections.extend(table_sections)

        image_section = self._extract_images(doc, file_descriptor)
        if image_section:
            sections.append(image_section)

        return ParsedDocument(document_id=file_descriptor.document_id, sections=sections)

    @staticmethod
    def _extract_tables(doc) -> list[ParsedSection]:
        table_sections: list[ParsedSection] = []
        for index, table in enumerate(doc.tables, start=1):
            row_lines: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    row_lines.append(" | ".join(cells))

            if row_lines:
                table_sections.append(
                    ParsedSection(
                        text="[TABLE]\n" + "\n".join(row_lines),
                        section_title=f"Table {index}",
                    )
                )
        return table_sections

    @staticmethod
    def _collect_alt_text(doc) -> dict[str, str]:
        """Map image relationship-id -> alt text/title (often names the diagram)."""
        alt_by_embed: dict[str, str] = {}
        for inline_shape in doc.inline_shapes:
            doc_pr = inline_shape._inline.docPr
            title = (doc_pr.get("title") or "").strip()
            descr = (doc_pr.get("descr") or "").strip()
            name = (doc_pr.get("name") or "").strip()
            label = descr or title or name
            try:
                embed = inline_shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            except Exception:
                embed = None
            if embed and label:
                alt_by_embed[embed] = label
        return alt_by_embed

    @classmethod
    def _extract_images(cls, doc, file_descriptor: FileDescriptor) -> ParsedSection | None:
        alt_by_embed = cls._collect_alt_text(doc)

        image_lines: list[str] = []
        visual_refs: list[dict[str, str]] = []
        index = 0

        for rel_id, rel in doc.part.rels.items():
            if "image" not in rel.reltype:
                continue
            index += 1
            label = alt_by_embed.get(rel_id, "")
            try:
                blob = rel.target_part.blob
                ext = rel.target_part.partname.rpartition(".")[2] or "png"
                name = label or f"image_{index}"
                path = write_image_asset(file_descriptor.file_path, blob, ext, f"img_{index}")
                visual_refs.append({"type": "image", "name": name, "path": path})
                detail = f" alt={label}" if label else ""
                image_lines.append(f"[IMAGE] name={name}{detail}")
            except Exception as error:
                logger.warning(
                    "DOCX image extraction failed",
                    file_path=str(file_descriptor.file_path),
                    error=str(error),
                )
                image_lines.append(f"[IMAGE] source={rel.target_ref}")

        if not image_lines:
            return None

        return ParsedSection(
            text="\n".join(dict.fromkeys(image_lines)),
            section_title="Images",
            metadata={"visual_refs": visual_refs},
        )
